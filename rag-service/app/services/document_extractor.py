import csv
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
from app.services.pdf_extractor import PDFExtractor


class DocumentExtractor:
    def __init__(self):
        self.pdf_extractor = PDFExtractor()

    def extract_text(self, file_path: Path, filename: str | None = None) -> str:
        extension = self._extension(filename or file_path.name)

        if extension == ".pdf":
            return self.pdf_extractor.extract_text(file_path)

        if extension == ".docx":
            return self._extract_docx(file_path)

        if extension in {".html", ".htm"}:
            return self._extract_html(file_path)

        if extension in {".xlsx", ".xls"}:
            return self._extract_excel(file_path)

        if extension == ".csv":
            return self._extract_csv(file_path)

        if extension in {".txt", ".md", ".markdown"}:
            return self._clean_text(file_path.read_text(encoding="utf-8", errors="ignore"))

        raise ValueError(f"Unsupported document type: {extension or 'unknown'}")

    def _extract_docx(self, file_path: Path) -> str:
        document = Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return self._clean_text("\n\n".join(parts))

    def _extract_html(self, file_path: Path) -> str:
        html = file_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        parts: list[str] = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th"]):
            text = tag.get_text(" ", strip=True)
            if text:
                parts.append(text)

        return self._clean_text("\n\n".join(parts))

    def _extract_excel(self, file_path: Path) -> str:
        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
        parts: list[str] = []

        for sheet_name, frame in sheets.items():
            frame = frame.fillna("")
            parts.append(f"Sheet: {sheet_name}")
            parts.append(frame.to_csv(index=False))

        return self._clean_text("\n\n".join(parts))

    def _extract_csv(self, file_path: Path) -> str:
        rows: list[str] = []

        with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                cleaned = [cell.strip() for cell in row if cell.strip()]
                if cleaned:
                    rows.append(" | ".join(cleaned))

        return self._clean_text("\n".join(rows))

    def _clean_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extension(self, filename: str) -> str:
        return Path(filename).suffix.lower()
